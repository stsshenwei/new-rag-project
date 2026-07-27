import tempfile
import unittest
from pathlib import Path

from app.services.document_parser import (
    BuiltinDocxParser,
    BuiltinExcelParser,
    BuiltinMarkdownParser,
    MarkdownFallbackParser,
    SUPPORTED_PARSE_EXTS,
    elements_from_markdown_and_html,
    get_parser_for_path,
)


class DocumentParserStructuredTests(unittest.TestCase):
    def test_get_parser_accepts_requested_types(self):
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            for name in ["a.pdf", "b.docx", "c.html", "d.xlsx", "e.md", "f.txt"]:
                path = base / name
                path.write_text("x", encoding="utf-8")
                self.assertIsNotNone(get_parser_for_path(path))

    def test_supported_parse_extensions_include_txt(self):
        self.assertIn(".txt", SUPPORTED_PARSE_EXTS)

    def test_builtin_parser_reads_txt_as_plain_text(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "manual.txt"
            path.write_text("DH-GPON-C+ 产品手册\n支持 GPON 上行和设备配置说明。", encoding="utf-8")

            parsed = BuiltinMarkdownParser().parse(path)

        self.assertEqual("txt", parsed.file_type)
        self.assertIn("DH-GPON-C+", parsed.markdown)
        self.assertTrue(any("GPON" in element.text for element in parsed.elements))

    def test_builtin_parser_reads_gb18030_txt(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "manual.txt"
            path.write_bytes("硬件参数\n光纤类型\n单模单纤".encode("gb18030"))

            parsed = BuiltinMarkdownParser().parse(path)

        self.assertIn("硬件参数", parsed.markdown)
        self.assertIn("光纤类型", parsed.markdown)
        self.assertIn("text_encoding:gb18030", parsed.diagnostics.warnings)

    def test_markdown_parser_normalizes_titles_paragraphs_and_tables(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "manual.md"
            path.write_text(
                "# 第4章\n\n## 资源规格\n\n下表列出服务器资源规格。\n\n"
                "| CPU | 内存 | GPU |\n|---|---|---|\n| Kunpeng 920 | 32*64GB | 8卡 |",
                encoding="utf-8",
            )

            parsed = MarkdownFallbackParser().parse(path)

        self.assertEqual("manual.md", parsed.file_name)
        self.assertEqual("md", parsed.file_type)
        self.assertTrue(any(element.type == "title" and element.title_path == "第4章/资源规格" for element in parsed.elements))
        tables = [element for element in parsed.elements if element.type == "table"]
        self.assertEqual(1, len(tables))
        self.assertIn("Kunpeng 920", tables[0].text)
        self.assertIn("CPU", tables[0].metadata["fields"])
        self.assertEqual("fallback", tables[0].metadata["parse_source"])
        self.assertIn("layout", tables[0].metadata)

    def test_builtin_markdown_parser_reports_builtin_source(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "manual.md"
            path.write_text("# Manual\n\nbody", encoding="utf-8")

            parsed = BuiltinMarkdownParser().parse(path)

        self.assertEqual("BuiltinMarkdownParser", parsed.diagnostics.parser_name)
        self.assertEqual("builtin_markdown", parsed.elements[0].metadata["parse_source"])

    def test_builtin_docx_parser_extracts_paragraphs_and_tables(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "manual.docx"
            doc = Document()
            doc.add_heading("Install", level=1)
            doc.add_paragraph("Follow the setup guide.")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "timeout"
            table.cell(1, 1).text = "30"
            doc.save(path)

            parsed = BuiltinDocxParser().parse(path)

        self.assertEqual("docx", parsed.file_type)
        self.assertIn("Install", parsed.markdown)
        self.assertTrue(any(element.metadata["parse_source"] == "builtin_docx" for element in parsed.elements))

    def test_builtin_excel_parser_extracts_sheet_markdown(self):
        from openpyxl import Workbook

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "inventory.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "设备"
            ws.append(["Name", "Count"])
            ws.append(["Switch", 2])
            wb.save(path)

            parsed = BuiltinExcelParser().parse(path)

        self.assertEqual("xlsx", parsed.file_type)
        self.assertIn("# Sheet: 设备", parsed.markdown)
        self.assertTrue(any(element.metadata["parse_source"] == "builtin_excel" for element in parsed.elements))

    def test_elements_preserve_caption_and_image_metadata(self):
        elements = elements_from_markdown_and_html(
            "# Manual\n\nTable 1: Options\n\n| Name | Value |\n|---|---|\n| timeout | 30 |\n\n![Architecture](images/arch.png)",
            "",
            parse_source="docling",
        )

        table = next(element for element in elements if element.type == "table")
        image = next(element for element in elements if element.type == "image")

        self.assertEqual("Table 1: Options", table.metadata["caption"])
        self.assertEqual([{"Name": "timeout", "Value": "30"}], table.metadata["rows"])
        self.assertEqual(1, table.metadata["row_count"])
        self.assertEqual(2, table.metadata["column_count"])
        self.assertEqual("docling", table.metadata["parse_source"])
        self.assertEqual("Architecture", image.metadata["caption"])
        self.assertEqual(["images/arch.png"], image.metadata["figure_refs"])


if __name__ == "__main__":
    unittest.main()
