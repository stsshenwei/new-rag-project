import csv
import html
import json
import re
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable

from docx import Document

SUPPORTED_EXTS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".doc",
    ".docx",
    ".pdf",
    ".html",
    ".htm",
    ".xlsx",
    ".xlsm",
    ".xls",
}

EXCLUDED_SOURCE_DIRS = {"processing_traces"}

DOC_TEXT_PATTERN = re.compile(
    r"[\u4e00-\u9fffA-Za-z0-9][\u4e00-\u9fffA-Za-z0-9\s\u3001\u3002\uff0c\uff01\uff1f\uff1b\uff1a\u201c\u201d\u2018\u2019\uff08\uff09\u3010\u3011\u300a\u300b,.!?:;\"'()\-_/]{7,}"
)


@dataclass(frozen=True)
class ParentChildChunk:
    source: str
    parent_id: str
    child_id: str
    parent_index: int
    child_index: int
    parent_text: str
    child_text: str
    section_title: str | None = None


class _VisibleTextHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self._skip_depth = 0
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, attrs):
        if tag.lower() in {"script", "style", "noscript"}:
            self._skip_depth += 1
        if tag.lower() in {"p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str):
        if tag.lower() in {"script", "style", "noscript"} and self._skip_depth > 0:
            self._skip_depth -= 1
        if tag.lower() in {"p", "div", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"}:
            self.parts.append("\n")

    def handle_data(self, data: str):
        if self._skip_depth:
            return
        text = html.unescape(data).strip()
        if text:
            self.parts.append(text)

    def text(self) -> str:
        lines = [" ".join(line.split()) for line in "".join(self.parts).splitlines()]
        return "\n".join(line for line in lines if line)


def iter_source_files(base_dir: Path) -> Iterable[Path]:
    for p in base_dir.rglob("*"):
        if p.name.startswith("~$"):
            continue
        try:
            relative_parts = set(p.relative_to(base_dir).parts[:-1])
        except ValueError:
            relative_parts = set()
        if relative_parts & EXCLUDED_SOURCE_DIRS:
            continue
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS:
            yield p


def read_text_with_fallback(path: Path) -> tuple[str, str]:
    content = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk", "big5"):
        try:
            return content.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace"), "utf-8-replace"


def load_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return read_text_with_fallback(path)[0]
    if suffix == ".csv":
        rows = []
        text, _ = read_text_with_fallback(path)
        reader = csv.reader(text.splitlines())
        for row in reader:
            rows.append(" | ".join(row))
        return "\n".join(rows)
    if suffix == ".json":
        text, _ = read_text_with_fallback(path)
        data = json.loads(text)
        return json.dumps(data, ensure_ascii=False, indent=2)
    if suffix == ".doc":
        return load_doc(path)
    if suffix == ".docx":
        return load_docx(path)
    if suffix == ".pdf":
        return load_pdf_as_markdown_with_pymupdf4llm(path)
    if suffix in {".html", ".htm"}:
        return load_html(path)
    if suffix in {".xlsx", ".xlsm", ".xls"}:
        return load_excel(path)
    return ""


def load_html(path: Path) -> str:
    parser = _VisibleTextHTMLParser()
    parser.feed(read_text_with_fallback(path)[0])
    parser.close()
    text = parser.text()
    if not text:
        raise ValueError(f"Failed to parse HTML file '{path.name}': no visible text extracted")
    return text


def load_excel(path: Path) -> str:
    try:
        import pandas as pd
    except Exception as exc:
        raise ValueError(
            "Excel parsing requires pandas plus an Excel engine such as openpyxl or xlrd. "
            "Please install backend requirements first."
        ) from exc

    try:
        sheets = pd.read_excel(path, sheet_name=None, dtype=str, keep_default_na=False)
    except Exception as exc:
        raise ValueError(f"Failed to parse Excel file '{path.name}': {exc}") from exc

    parts: list[str] = []
    for sheet_name, frame in sheets.items():
        parts.append(f"# Sheet: {sheet_name}")
        if frame.empty:
            continue
        headers = [str(col).strip() for col in frame.columns]
        if any(headers):
            parts.append(" | ".join(headers))
        for row in frame.itertuples(index=False, name=None):
            values = [str(value).strip() for value in row]
            row_text = " | ".join(value for value in values if value)
            if row_text:
                parts.append(row_text)

    text = "\n".join(part for part in parts if part).strip()
    if not text:
        raise ValueError(f"Failed to parse Excel file '{path.name}': no readable cells extracted")
    return text


def load_doc(path: Path) -> str:
    raw = path.read_bytes()
    parts = [
        _extract_doc_text(raw.decode("utf-16le", errors="ignore")),
        _extract_doc_text(raw.decode("latin1", errors="ignore")),
    ]
    text = "\n".join(part for part in parts if part).strip()
    if not text:
        raise ValueError(
            f"Failed to parse DOC file '{path.name}': no readable text could be extracted"
        )
    return text


def _extract_doc_text(text: str) -> str:
    cleaned: list[str] = []
    seen: set[str] = set()

    for match in DOC_TEXT_PATTERN.findall(text):
        line = " ".join(match.split()).strip()
        if len(line) < 8 or line in seen:
            continue
        seen.add(line)
        cleaned.append(line)

    return "\n".join(cleaned)


def load_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(f"Failed to parse DOCX file '{path.name}': {exc}") from exc

    parts: list[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join([c for c in cells if c])
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def load_pdf_as_markdown_with_pymupdf4llm(path: Path) -> str:
    try:
        import pymupdf4llm
    except Exception as exc:
        raise ValueError(
            "PDF parsing requires pymupdf4llm and pymupdf. "
            "Please install backend requirements first."
        ) from exc

    try:
        # Disable layout model pipeline to avoid ONNX runtime dependency issues on Windows.
        if hasattr(pymupdf4llm, "use_layout"):
            pymupdf4llm.use_layout(False)
        markdown = pymupdf4llm.to_markdown(str(path))
    except Exception as exc:
        raise ValueError(f"Failed to parse PDF file '{path.name}' with PyMuPDF4LLM: {exc}") from exc

    markdown = (markdown or "").strip()
    if not markdown:
        raise ValueError(f"Failed to parse PDF file '{path.name}': no readable markdown extracted")
    return markdown


def split_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    text = " ".join(text.split())
    if not text:
        return []

    chunks = []
    start = 0
    step = max(chunk_size - overlap, 1)

    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += step

    return chunks


def split_markdown_by_headers(markdown_text: str, chunk_size: int, overlap: int) -> list[str]:
    try:
        from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter
    except Exception as exc:
        raise ValueError(
            "Markdown header splitting requires langchain-text-splitters. "
            "Please install backend requirements first."
        ) from exc

    headers_to_split_on = [
        ("#", "h1"),
        ("##", "h2"),
        ("###", "h3"),
        ("####", "h4"),
    ]
    header_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on, strip_headers=False)
    docs = header_splitter.split_text(markdown_text)

    # Secondary split keeps chunk size stable while preserving section boundaries.
    rc_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    final_docs = rc_splitter.split_documents(docs)
    chunks = [d.page_content.strip() for d in final_docs if d.page_content and d.page_content.strip()]
    if chunks:
        return chunks
    return split_text(markdown_text, chunk_size, overlap)


def split_by_file_type(path: Path, text: str, chunk_size: int, overlap: int) -> list[str]:
    suffix = path.suffix.lower()
    if suffix in {".pdf", ".md", ".markdown", ".html", ".htm"}:
        return split_markdown_by_headers(text, chunk_size, overlap)
    return split_text(text, chunk_size, overlap)


def build_parent_child_chunks(
    source: str,
    text: str,
    parent_size: int,
    parent_overlap: int,
    child_size: int,
    child_overlap: int,
) -> list[ParentChildChunk]:
    parent_chunks = split_text(text, parent_size, parent_overlap)
    items: list[ParentChildChunk] = []

    for parent_index, parent_text in enumerate(parent_chunks):
        parent_id = f"{source}::parent-{parent_index}"
        child_chunks = split_text(parent_text, child_size, child_overlap)
        for child_index, child_text in enumerate(child_chunks):
            child_id = f"{parent_id}::child-{child_index}"
            items.append(
                ParentChildChunk(
                    source=source,
                    parent_id=parent_id,
                    child_id=child_id,
                    parent_index=parent_index,
                    child_index=child_index,
                    parent_text=parent_text,
                    child_text=child_text,
                )
            )

    return items
